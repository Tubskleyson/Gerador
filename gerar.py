import sys
import os
from win32com import client
from docx import Document

def get_run(modelo, n):

    print(n)

    runs = [0 for i in range(n)]

    for i in modelo.tables:
        x = len(i.columns)
        y = len(i.rows)

        for j in range(x):
            for k in range(y):
                ps = i.cell(k, j).paragraphs
                for l in ps:
                    for z in l.runs:
                        #print(z.text)

                        for w in range(n):
                            if z.text == 'Valor_'+str(w+1): runs[w] = z

                        #print(runs)
                        if all(y != 0 for y in runs): return runs

wdFormatPDF = 17

print('\n> Modelo do Certificado')
file_path = os.path.abspath(input('\n< '))

print('\n> Lista de valores')
list_path = os.path.abspath(input('\n< '))

print('\n> Pasta de sáida')
out_path = os.path.abspath(input('\n< '))

if '\\' != out_path[-1]: out_path+='\\'

nomes = sorted([i.split(';') for i in open(list_path).readlines()[1:]])[1:]


for i in range(len(nomes)):
    nomes[i] = [j.strip('\n') for j in nomes[i]]



modelo = Document(file_path)


runs = get_run(modelo,len(nomes[0]))

word = client.Dispatch('Word.Application')


j = 1

for i in nomes:

    print('> Gerano arquivo %d de %d' %(j,len(nomes)),end='\r')
    j+=1

    for k in range(len(i)):
        runs[k].text = i[k]

    modelo.save(out_path+i[0]+'.docx')

    doc = word.Documents.Open(os.path.abspath(out_path+i[0] + '.docx'))
    doc.SaveAs(out_path+ i[0] + '.pdf', FileFormat=wdFormatPDF)
    doc.Close()

    os.remove(out_path+i[0]+'.docx')


word.Quit()
