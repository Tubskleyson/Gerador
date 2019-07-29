#!/usr/bin/python3

from subprocess import  Popen, PIPE
from docx import Document
from sys import argv
import os



def get_runs(doc : Document, nomes : list):

    n = len(nomes[0])

    runs = [0 for i in range(n)]

    for i in doc.tables:
        x = len(i.columns)
        y = len(i.rows)

        for j in range(x):
            for k in range(y):
                ps = i.cell(k, j).paragraphs
                for l in ps:
                    for z in l.runs:

                        for w in range(n):
                            if z.text == 'Valor_' + str(w + 1): runs[w] = z

                        if all(y != 0 for y in runs): return runs

    for j in doc.paragraphs:


        for z in j.runs:

            for w in range(n):
                if z.text == 'Valor_' + str(w + 1): runs[w] = z

            if all(y != 0 for y in runs): return runs

    return []





n_argumentos = len(argv)

if n_argumentos < 3:

    print("[ ERRO ]  Há argumentos faltando")
    print("Uso : gerarcertificados [ modelo.docx ] [ lista.csv ] [ diretorio de saida ]")

    exit()


modelo = os.path.abspath(argv[1])

if modelo.split('.')[-1] != 'docx':

    print("[ ERRO ]  Arquivo modelo deve ser do tipo .docx")
    exit()

if not os.path.exists(modelo):

    print("[ ERRO ]  Não foi possível encontrar o arquivo %s" %modelo)
    exit()


lista = os.path.abspath(argv[2])

if lista.split('.')[-1] != 'csv':

    print("[ ERRO ]  Arquivo de lista deve ser do tipo .csv")
    exit()


if not os.path.exists(lista):

    print("[ ERRO ]  Não foi possível encontrar o arquivo %s" %lista)
    exit()

output = os.path.abspath('.')

if n_argumentos > 3:

    output = os.path.abspath(argv[3])

    if not os.path.exists(output):

        print("[ ERRO ]  Não foi possível encontrar a pasta %s" %output)
        exit()



nomes = sorted([i.split(';') for i in open(lista).readlines()])

n_nomes = len(nomes)

for i in range(len(nomes)):

    nomes[i] = [j.strip('\n') for j in nomes[i]]


doc = Document(modelo)

runs = get_runs(doc, nomes)

print("\nAnalisando os dados \n")


if not runs:

    print("[ ERRO ]  Não foi possível encontrar os campos a serem preenchidos")
    exit()

print("", end='\r')

c = 1

for info in nomes:

    hashtags = int(40*c/n_nomes)

    print("< Gerando arquivos {}{} {}/{}".format(chr(9608)*hashtags, ' '*(40-hashtags), c, n_nomes), end = '\r')

    for index in range(len(info)):

        runs[index].text = info[index]


    while os.path.exists(os.path.join(output, info[0].replace(' ', '_') + '.pdf')):

        info[0] += '_'


    docx_file = os.path.join(output, info[0] + '.docx')

    doc.save(docx_file)

    p1 = Popen(["soffice", "--headless", "--convert-to", "pdf", docx_file, "--outdir", output], stdout=PIPE)
    p1.communicate()

    os.remove(docx_file)

    c += 1

print("\n")