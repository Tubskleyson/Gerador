from tkinter import filedialog,ttk
from tkinter import *
from PIL import ImageTk, Image
from win32com import client
from docx import Document
import os

class App:

    def __init__(self):

        self.window = Tk()

        self.window.iconbitmap('assets/certi.ico')

        self.window.title('Certificados')

        self.window.resizable(False,False)

        self.window.geometry('800x500')

        self.window.configure(background='white')

        self.lista = ''
        self.modelo = ''
        self.output = ''
        self.pasta = ''

        self.nomes = []


        space = Label(self.window,bg='white')
        space.pack(pady=20)


        load = Image.open("assets/cert.png")

        load = load.resize((200, 200), Image.ANTIALIAS)

        img = ImageTk.PhotoImage(load)



        image = Label(self.window,text = 'kalfoi', image = img, background='white')
        image.image = img
        image.pack()

        tit = Label(self.window,text='Gerador de Certificados', bg='white',font="sans 25 ", fg='grey')
        tit.pack()

        bt = Button(self.window,text='Novo Pacote', height=2, width=13, cursor='hand2', command=self.select)
        bt.pack(pady=30)

        cred = Label(self.window, text='[Tubs]', bg='white', font="system 9 bold ", fg='lightgrey')
        cred.pack(side=BOTTOM,pady=5)

        self.widgets = [space,image,tit,bt]

    def select(self):

        for i in self.widgets: i.destroy()

        self.tit = Label(self.window, text='Selecionar Documentos', bg='white', font="sans 25 ", fg='grey')
        self.tit.pack(pady=30)

        self.mid = Frame(bg='white')
        self.mid.pack(fill=X)


        left = Frame(self.mid,bg='white')
        left.pack(side=LEFT, fill=Y,pady=50)

        lp = Image.open("assets/cert.jpg")

        lp = lp.resize((80, 80), Image.ANTIALIAS)

        li = ImageTk.PhotoImage(lp)

        imagel = Label(left, text='kalfoi', image=li, background='white')
        imagel.image = li
        imagel.pack(side=TOP)



        b0 = Button(left,cursor='hand2', text='Selecionar Modelo (.docx)', width=25,height=2, command=lambda : self.get_file(0))
        b0.pack(padx=100, pady=20)

        self.lbl0 = Label(left, text="Nenhum arquivo selecionado", bg='white')
        self.lbl0.pack()

        right = Frame(self.mid,bg='white')
        right.pack(side=RIGHT, fill=Y, pady=50)

        rp = Image.open("assets/list.png")

        rp = rp.resize((80, 80), Image.ANTIALIAS)

        ri = ImageTk.PhotoImage(rp)

        imager = Label(right, text='kalfoi', image=ri, background='white')
        imager.image = ri
        imager.pack(side=TOP)

        b1 = Button(right,cursor='hand2', text='Selecionar Lista (.csv)', width=25,height=2, command=lambda : self.get_file(1))
        b1.pack(padx=100, pady=20)

        self.lbl1 = Label(right, text="Nenhum arquivo selecionado", bg='white')
        self.lbl1.pack()

        self.bot = Frame(bg='white')
        self.bot.pack(fill=BOTH,pady=10)


        self.continuar = Button(self.bot,text = 'Continuar',bg='green',height=2,width=15,fg='white',activebackground='green', activeforeground='white',font="sans 13 bold", command=self.check)
        self.continuar.pack()

        self.widgets = [left, right]

    def get_file(self,x):

        if not x:
            self.modelo =  filedialog.askopenfilename(initialdir="/", title="Select file",filetypes=(("Documento Word", "*.docx"), ("all files", "*.*")))
            if self.modelo: self.lbl0['text'] = self.modelo.split('/')[-1]
            else: self.lbl0['text'] = 'Nenhum arquivo selecionado'

        elif x==1:
            self.lista =  filedialog.askopenfilename(initialdir="/", title="Select file",filetypes=(("Planilha", "*.csv"), ("all files", "*.*")))
            if self.lista: self.lbl1['text'] = self.lista.split('/')[-1]
            else: self.lbl1['text'] = 'Nenhum arquivo selecionado'

        elif x==2:
            self.pasta = filedialog.askdirectory(initialdir="/",title ='Selecione a pasta')+'/'
            if self.pasta: self.lbl0['text'] = self.pasta
            else: self.lbl0['text'] = 'Nenhum pasta selecionada'

    def get_runs(self):

        n = len(self.nomes[0])

        runs = [0 for i in range(n)]

        for i in self.doc.tables:
            x = len(i.columns)
            y = len(i.rows)


            for j in range(x):
                for k in range(y):
                    ps = i.cell(k, j).paragraphs
                    for l in ps:
                        for z in l.runs:

                            for w in range(n):
                                if z.text == 'Valor_' + str(w + 1): runs[w] = z

                            if all(y != 0 for y in runs): return runs

        for j in self.doc.paragraphs:

            print(j.runs)

            for z in j.runs:

                for w in range(n):
                     if z.text == 'Valor_' + str(w + 1): runs[w] = z

                if all(y != 0 for y in runs): return runs

        return []

    def check(self):

        erro = 0

        self.window['cursor'] = 'arrow'

        if not self.lista: erro = 'Você esqueceu de adicionar a lista'
        elif not self.modelo: erro = 'Você esqueceu de adicionar o modelo'
        elif not self.nomes:


            self.nomes = sorted([i.split(';') for i in open(self.lista).readlines()])

            for i in range(len(self.nomes)):
                self.nomes[i] = [j.strip('\n') for j in self.nomes[i]]

            self.doc = Document(self.modelo)

            self.runs = self.get_runs()

            if not self.runs: erro = 'Não foi possível encontrar os campos \na serem preenchidos'

        elif not self.pasta: erro = 'Você não selecionou nenhuma pasta'


        if erro:

            popup = Toplevel(bg='white')
            popup.title('Erro')
            popup.geometry("300x150")

            aviso = Label(popup,text=erro,bg='white')
            aviso.pack(pady=30)

            bt = Button(popup,width=6,height=2, text='Ok', command=popup.destroy)
            bt.pack()

            return


        if not self.pasta: self.get_pasta()
        else: self.gerar()

    def check_load(self):

        self.window['cursor'] = 'watch'
        self.window.after(10,self.check)

    def get_pasta(self):

        for i in self.widgets: i.destroy()

        self.tit['text'] = 'Selecionar Pasta de Destino'

        center = Frame(self.mid, bg='white')
        center.pack(fill=Y, pady=50)

        lp = Image.open("assets/pasta.png")

        lp = lp.resize((80, 80), Image.ANTIALIAS)

        li = ImageTk.PhotoImage(lp)

        imagel = Label(center, text='kalfoi', image=li, background='white')
        imagel.image = li
        imagel.pack(side=TOP)

        b0 = Button(center, cursor='hand2', text='Selecionar Pasta', width=25, height=2,
                    command=lambda: self.get_file(2))
        b0.pack(padx=100, pady=20)

        self.lbl0 = Label(center, text="Nenhuma pasta selecionada", bg='white')
        self.lbl0.pack()

        self.continuar['text'] = 'Gerar'
        self.continuar['command'] = self.check_load

        self.widgets = [center]

    def geraUm(self):

        self.window['cursor'] = 'arrow'


        i = self.nomes[self.j]


        self.bar.step((1 / len(self.nomes))*100)
        self.lbl0['text'] = "Gerando arquivo %d / %d" % (self.j+1, len(self.nomes))

        self.j += 1



        for k in range(len(i)):
            self.runs[k].text = i[k]


        while os.path.exists(self.pasta + i[0].replace(' ','_') + '.pdf'):

            i[0] += '_'


        self.doc.save(self.pasta + i[0] + '.docx')

        doc = self.word.Documents.Open(os.path.abspath(self.pasta + i[0] + '.docx'))
        doc.SaveAs(os.path.abspath(self.pasta + i[0].replace(' ','_') + '.pdf'), FileFormat=17)
        doc.Close()

        os.remove(self.pasta + i[0] + '.docx')

    

        if self.j<len(self.nomes): self.window.after(100,self.geraUm)
        else:
            self.window['cursor'] = 'watch'
            self.window.after(100,self.finish)


    def gerar(self):

        for i in self.widgets: i.destroy()

        self.continuar.destroy()

        self.tit['text'] = 'Gerando Arquivos'

        center = Frame(self.mid, bg='white')
        center.pack(fill=Y, pady=70)

        self.lbl0 = Label(center, text="Gerando arquivo 0 / %d" %len(self.nomes), bg='white',font="sans 13 ")
        self.lbl0.pack()

        self.bar = ttk.Progressbar(center,orient='horizontal',length=500)
        self.bar.pack(fill=X,pady=40)

        self.window['cursor'] = 'watch'

        self.word = client.Dispatch('Word.Application')

        self.j = 0
        self.window.after(100,self.geraUm)


    def finish(self):

        self.lbl0['text'] = "Operação Concluída"

        self.continuar = Button(self.bot, text='Finalizar', bg='green', height=2, width=15, fg='white',
                                activebackground='green', activeforeground='white', font="sans 13 bold",
                                command=self.window.destroy)
        self.continuar.pack()

        self.bar.destroy()

        self.word.Quit()

        self.window['cursor'] = 'arrow'








app = App()
app.window.mainloop()
